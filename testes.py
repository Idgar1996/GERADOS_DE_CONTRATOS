import customtkinter as ctk
from customtkinter import *
from PIL import Image
from customtkinter import CTkComboBox
import openpyxl
from openpyxl import Workbook
import pandas as pd
from tkinter.filedialog import askopenfilename
from docx import Document
from tkinter.filedialog import asksaveasfilename

# --------------------------------------------------------------------------------
# ===> Inicialização da Aplicação <===
app = CTk() # Define que o 'app' será executado (aberto) utilizando o customtkinter
app.title('Gerador de Contratos') # Define o título a ser exibido na 'janela' do app
app.resizable(True, False) # Determina se a 'janela' será redimensionavel ou não sendo cada um dos parâmetros: (largura, altura)
app._set_appearance_mode('dark') # Define o tema padrão ao abrir o aplicativo


# --------------------------------------------------------------------------------
# ===> Centralização da Janela <===
largura_app = 1440 # Define a largura ja 'janela'
altura_app = 900 # Define a 'altura' da janela
largura_tela = app.winfo_screenwidth() # Informa o tamanho da tela (do monitor)
altura_tela = app.winfo_screenheight() # Informa a altura da tela (do monitor)
x = int((largura_tela - largura_app) / 2) # Define a posição em relacão ao eixo X, ou seja, na horizontal
y = int((altura_tela - altura_app) / 2) # Define a posição em relação ao eixo Y, ou seja, vertical
app.geometry(f"{largura_app}x{altura_app}+{x}+{y}") # Define a posição em que a 'janela' será aberta, tendo como referência o lado esquerdo, ou seja: largura_app = largura da 'janela'; altura_app = altura da 'janela'; x e y são as posições com os eixos x partindo da esquerda e y partindo do topo


# --------------------------------------------------------------------------------
# ===> Funções <===

# ==> Função para mostrar o frame seguinte <==
def mostrar_frame(frame): # Função que define uma variável para um especifico parâmetro, seja ela botão, frame, label; 'def' sempre deve ser usada para DEFinir uma função
    frame.tkraise() # Define que a função 'mostrar_frame', quando chamada fará com que um determinado frame seja exibido em frente aos outros, esse frame sendo especificado entre os parênteses (pg1 documento comandos)



# ==> Função para fazer o login <==
def login_realizado(): # Define todos os parâmetro para que a função seja executada corretamente futuramente, seja por um botão algum outro parametro
    global cargo_usuario_logado  # Quando é feito o contrato, o cliente decide qual cargo ele terá, e essa informação que será usada/modificada dentro dessa função especifica

    nome = entry_nome.get().strip() # Define que a variável 'nome' será o nome que o usuário digitar no 'entry_nome; o '.get' pega essa informação ao cliente digita-lá no campor especificado; '.strip' apaga os espaços no começo e final da string.
    senha = entry_senha_login.get().strip()# Define que a variável 'senha' será o nome que o usuário digitar no 'entry_senha; o '.get' pega essa informação ao cliente digita-lá no campor especificado; '.strip' apaga os espaços no começo e final da string.

    # Verifica se os campos estão preenchidos
    if not nome or not senha: # Define caso algum dado não tenha sido preenchido ou não nos entrys da página de login
        resultado_login.configure(text="Preencha todos os campos!") # Exibe a mensagem especificada para caso um dos, ou nenhum dos dois entrys especificados anteriormente tenha sido 'satisfeitos, 'preenchido', que será exibida em um Label chamado resultado
        app.after(3000, lambda: resultado_login.configure(text=''))
        return # Após a verificação anterior, 'volta' para a base da função para executar novas tarefas

    try: # Ele tenta executar um bloco de código, e se der erro, ele pula para uma outra parte chamada except onde você pode dizer o que fazer quando isso acontece; Ele evita quebrar seu código.
        arquivo_excel = "usuarios_registrados.xlsx" # Define que a variavel 'arquivo_excel' será um determinado arquivo no computador nesse formato; Sempre definir na variavel o tipo e garantir que o arquivo tenha a extensão correta no fim do nome
        df = pd.read_excel(arquivo_excel) # Variavel que solicita a biblioteca 'pandas' para que leia e interprete o arquivo ja determinado anteriormente

        if "Senha" in df.columns: # Determina se há uma coluna (.columns) que tenha na primeira célula a palavra 'Senha'
            df["Senha"] = df["Senha"].astype(str) # Caso tenha uma coluna com esse nome, determina que todos itens daquela coluna serão convertidos (.astype) em string, independente do tipo
            usuario = df[(df["Nome"] == nome) & (df["Senha"] == senha)] # Variavel temporária que verfica se alguma palavra dentro das colunas especificadas ('Nome' e 'Senha') correspondem com os dados inseridos pelo usuario nos campos entry da pagina de login; O '&' é um valor lógico, ou seja, os dois dados precisam estar iguais
        if not usuario.empty: # Verifica se a variavel definida anteriormente ('usuario' que são os nome e senha digitados) não está vazia('empty') ela retorna um valor positivo para seguir para a proxima tarefa
            cargo_usuario_logado = usuario["Cargo"].values[0]  # Le a o valor da coluna 'cargo', na linha do nome do usuario digitado anteriormente, e armazena o primeiro valor, se o cargo for "Gerente", agora a variável cargo_usuario_logado guarda isso.
            resultado_login.configure(text="") # Limpa o campo 'resultado', caso tenha sido exibido alguma mensagem anteriormente; 'configure()' é usado para atualizar o texto dos labels
            mostrar_frame(frame_opcoes)  # Redireciona para a tela de opções
        else: # Caso nenhum dos parametros acima tenham sido satisfeitos, executa o a funcão abaixo
            resultado_login.configure(text="Nome ou senha incorretos!", text_color="#FFD700") # Define o texto que sera exibido caso, os parametros 'nome' e 'senha' tenham sido preenchidos do forma diferente do armazenado no arquivo excel
            app.after(3000, lambda: resultado_login.configure(text=''))

    except FileNotFoundError: # Função que é usada quando um determinado arquivo, não é localizado, nesse caso o documento excel
        resultado_login.configure(text="Arquivo de usuários não encontrado!", text_color="#FFD700") # Mensagem definida para quando um determinado arquivo já determinado, não seja mais localizado, que será exibida em um Label chamado resultado
    except Exception as e: # Define a variavel 'e' para a condição, caso seja um erro diferente de algum ja especificado anteriormente, no caso, o erro de arquivo nao encontrado
        resultado_login.configure(text=f"Erro: {e}", text_color="red") # Mensagem definida para quando ocorrer um erro diferente do erro de nao localizar algum arquivo, que será exibida em um Label chamado resultado

# ==> Função para alternar a visibilidade das senhas na tela de registro <==
def alternar_visibilidade_senha_registro(): # Função que será usada para habilitar ou desabilitar a visibilidade das senhas
    if checkbox_mostrar_senha_registro.get() == 1: # Define que a se a caixa de seleção estiver desmarcada (0 = false {caso a checkbox esteja descelecionado}, 1 = true {caso a checkbox esteja selecionada})
        entry_senha_registro.configure(show="") # Não exibe nada no lugar da informação inserida pelo usuário no campo 'entry_senha_registro'
        entry_confirma_senha_registro.configure(show="") # Não exibe nada no lugar da informação inserida pelo usuário no campo 'entry_confirma_senha_registro'
    else: # Função que define para a situação inversa, caso a checkbox não tenha sido marcada pelo cliente
        entry_senha_registro.configure(show="*") # Exibe no lugar da informação inserida pelo cliente no campo 'entry_senha_registro', o caracter '*', com intuito de esconder a senha do cliente
        entry_confirma_senha_registro.configure(show="*") # Exibe no lugar da informação inserida pelo cliente no campo 'entry_confirma_senha_registro', o caracter '*', com intuito de esconder a senha do cliente


# ==> Função para alternar a visibilidade das senhas na tela de recuperação <==
def alternar_visibilidade_senha_recuperar(): # Função que será usada para habilitar ou desabilitar a visibilidade das senhas
    if checkbox_mostrar_senha_recuperar.get() == 1: # Define que a se a caixa de seleção estiver desmarcada (0 = false {caso a checkbox esteja descelecionado}, 1 = true {caso a checkbox esteja selecionada})
        entry_nova_senha.configure(show="") # Não exibe nada no lugar da informação inserida pelo usuário no campo 'entry_nova_senha'
        entry_confirma_nova_senha.configure(show="") # Não exibe nada no lugar da informação inserida pelo usuário no campo 'entry_confirma_nova_senha'
    else: # Função que define para a situação inversa, caso a checkbox não tenha sido marcada pelo cliente
        entry_nova_senha.configure(show="*") # Exibe no lugar da informação inserida pelo cliente no campo 'entry_nova_registro', o caracter '*', com intuito de esconder a senha do cliente
        entry_confirma_nova_senha.configure(show="*") # Exibe no lugar da informação inserida pelo cliente no campo 'entry_confirma_nova_senha', o caracter '*', com intuito de esconder a senha do cliente

# ==> Função para alternar a visibilidade das senhas na tela de login <==
def alternar_visibilidade_senha_login(): # Função que é usada para habilitar ou desabilitar a visibilidade da senha
    if checkbox_mostrar_senha_login.get() == 1: # Define que a se a caixa de seleção estiver desmarcada (0 = false {caso a checkbox esteja descelecionado}, 1 = true {caso a checkbox esteja selecionada})
        entry_senha_login.configure(show='') # Não exibe nada no lugar da informação inserida pelo usuário no campo 'entry_senha_login'
    else: # Função que define para a situação inversa, caso a checkbox não tenha sido marcada pelo cliente
        entry_senha_login.configure(show='*') # Exibe no lugar da informação inserida pelo cliente no campo 'entry_senha_login', o caracter '*', com intuito de esconder a senha do cliente

# ==> Função para registrar o usuário <==
def alertas_registrar_usuario(): # Função que será usada para efetuar o registro dos dados inseridos pelo usuário

    mensagem_registro.configure(text="") # Limpa a mensagem definida pela função caso o usuario insira dois dados diferentes no campo de entry_senha e entry_confirma_senha

    nome = entry_nome_registro.get().strip() # Variavel temporária para pegar o nome inserido pelo usuario no campo entry_nome_registro
    telefone = entry_telefone_registro.get().strip() # Variavel temporária para pegar o telefone inserido pelo usuario no campo entry_telefone_registro
    cargo = combobox_cargo_registro.get().strip() # Variavel temporária para pegar o cargo selecionado no combobox_cargo_registro
    senha = entry_senha_registro.get().strip() # Variavel temporária para pegar a senha inserida pelo usuario no campo entry_senha_registro
    confirma_senha = entry_confirma_senha_registro.get().strip() # Variavel temporária para pegar senha inserida reinserida pelo usuario no campo entry_confirma_senha_registro

    pergunta1 = combobox_pergunta1.get().strip() # Variavel temporária que pega a pergunta 1 escolhida pelo usuario no combobox_pergunta1
    resposta1 = entry_resposta1.get().strip() # Variavel temporária para pegar a resposta inserida pelo usuario no campo entry_resposta1
    pergunta2 = combobox_pergunta3.get().strip() # Variavel temporária que pega a pergunta 1 escolhida pelo usuario no combobox_pergunta2
    resposta2 = entry_resposta2.get().strip() # Variavel temporária para pegar a resposta inserida pelo usuario no campo entry_resposta2
    pergunta3 = combobox_pergunta3.get().strip() # Variavel temporária que pega a pergunta 1 escolhida pelo usuario no combobox_pergunta3
    resposta3 = entry_resposta3.get().strip() # Variavel temporária para pegar a resposta inserida pelo usuario no campo entry_resposta3

    if not nome or not telefone or not cargo or not senha or not confirma_senha: # Verifica se todos os campos na tela de registro foram preenchidos
        mensagem_registro.configure(text="Preencha todos os campos!", text_color="red") # Caso algum dado não tenha sido preenchido, exibe o aviso para que o cliente insira todos os dados
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça
        return # Retorna o resultado solicitado na função acima, executa a ação que foi solicitada, no caso, exibir a mensagem caso o cliente não tenha inserido algum dado nos campos

    if cargo == "Selecione": # Função que define se algum cargo foi selecionado
        mensagem_registro.configure(text="Selecione um cargo válido!", text_color="red") # Caso algum cargo não tenha sido selecionado, exibe o aviso para que o cliente selecione um cargo
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça
        return # Retorna o resultado solicitado na função acima, executa a ação que foi solicitada, no caso, exibir a mensagem caso o cliente não tenha inserido algum dado nos campos

    if senha != confirma_senha: # Função que define se a senha inserida é igual a senha
        mensagem_registro.configure(text="As senhas não coincidem!", text_color="red") # Caso as senhas inseridas não coincidam, exibe o aviso para que o cliente digite as senhas iguais
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça
        return # Retorna o resultado solicitado na função acima, executa a ação que foi solicitada, no caso, exibir a mensagem caso o cliente não tenha inserido algum dado nos campos

    if pergunta1 == "Selecione" or pergunta2 == "Selecione" or pergunta3 == "Selecione": # Verifica se as perguntas de segurança foram selecionada
        mensagem_registro.configure(text="Selecione todas as perguntas de segurança!", text_color="red") # Caso algum cargo não tenha sido selecionado, exibe a mensagem para que as perguntas de segurança sejam selecionadas
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça
        return # Retorna o resultado solicitado na função acima, executa a ação que foi solicitada, no caso, exibir a mensagem caso o cliente não tenha inserido algum dado nos campos
    
    if not resposta1 or not resposta2 or not resposta3: # Verifica se todos os campos na tela de registro foram preenchidos
        mensagem_registro.configure(text='Responda todas perguntas de segurança!', text_color='red') # Caso algum dado não tenha sido preenchido, exibe o aviso para que o cliente insira todos os dados
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça
        return # Retorna o resultado solicitado na função acima, executa a ação que foi solicitada, no caso, exibir a mensagem caso o cliente não tenha inserido algum dado nos campos

    try:
        # Salva os dados no Excel
        salvar_dados_no_excel_registro(nome, telefone, cargo, senha, pergunta1, resposta1, pergunta2, resposta2, pergunta3, resposta3)
        mensagem_registro.configure(text="Cadastro realizado com sucesso!", text_color="green")
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça

    except Exception as e:
        mensagem_registro.configure(text=f"Erro ao salvar: {e}", text_color="red")
        app.after(3000, lambda: mensagem_registro.configure(text='')) # Define um tempo especifico para que a mensagem exibida desapareça


# ==> Função para salvar os dados no Excel <==
def salvar_dados_no_excel_registro(nome, telefone, cargo, senha, pergunta1, resposta1, pergunta2, resposta2, pergunta3, resposta3):
    arquivo_excel = "usuarios_registrados.xlsx"

    try:
        # Tenta abrir o arquivo existente
        workbook = openpyxl.load_workbook(arquivo_excel)
        sheet = workbook.active
    except FileNotFoundError:
        # Se o arquivo não existir, cria um novo
        workbook = Workbook()
        sheet = workbook.active
        # Adiciona cabeçalhos
        sheet.append(["Nome", "Telefone", "Cargo", "Senha", "Pergunta 1", "Resposta 1", "Pergunta 2", "Resposta 2", "Pergunta 3", "Resposta 3"])

    # Adiciona os dados do usuário
    sheet.append([nome, telefone, cargo, senha, pergunta1, resposta1, pergunta2, resposta2, pergunta3, resposta3])

    # Salva o arquivo
    workbook.save(arquivo_excel)

# ==> Funcao para limpar os campos da tela de login <==
def limpar_campos_login():
    entry_nome.delete(0, 'end')
    entry_senha_login.delete(0, 'end')
    checkbox_mostrar_senha_login.deselect()  # Desmarca o checkbox "Mostrar Senha"

# ==> Função para limpar os campos e ocultar a mensagem <==
def limpar_campos_registro():
    # Limpa os campos da tela de registro
    entry_nome_registro.delete(0, 'end')
    entry_telefone_registro.delete(0, 'end')
    combobox_cargo_registro.set("Selecione")
    entry_senha_registro.delete(0, 'end')
    entry_confirma_senha_registro.delete(0, 'end')

    # Limpa os campos das perguntas de segurança
    combobox_pergunta1.set("Selecione")
    entry_resposta1.delete(0, 'end')
    combobox_pergunta2.set("Selecione")
    entry_resposta2.delete(0, 'end')
    combobox_pergunta3.set("Selecione")
    entry_resposta3.delete(0, 'end')

    # Desmarca o checkbox "Mostrar Senha"
    checkbox_mostrar_senha_registro.deselect()

# ==> Função para limpar os campos da tela de recuperação de senha <==
def limpar_campos_recuperacao():
    entry_seu_nome.delete(0, 'end')
    entry_nova_senha.delete(0, 'end')
    entry_confirma_nova_senha.delete(0, 'end')
    combobox_pergunta_recuperacao1.set("Selecione")
    entry_resposta_recuperacao1.delete(0, 'end')
    combobox_pergunta_recuperacao2.set("Selecione")
    entry_resposta_recuperacao2.delete(0, 'end')
    combobox_pergunta_recuperacao3.set("Selecione")
    entry_resposta_recuperacao3.delete(0, 'end')


# ==> funcao para limpar os campos da tela de dados do cliente <==
def limpar_campos_dados_cliente():
    # Limpa os campos de dados pessoais
    entry_nome_cliente.delete(0, 'end')
    entry_cpf_cnpj_cliente.delete(0, 'end')
    entry_rua_cliente.delete(0, 'end')
    entry_numero_cliente.delete(0, 'end')
    entry_cidade_cliente.delete(0, 'end')
    entry_uf_cliente.delete(0, 'end')
    entry_cep_cliente.delete(0, 'end')

    # Limpa os campos de dados de negociação
    entry_qnts_modulos.delete(0, 'end')
    entry_tipo_modulos.delete(0, 'end')
    entry_inversores.delete(0, 'end')
    entry_potencia_inversores.delete(0, 'end')
    entry_endereco_instalacao.delete(0, 'end')
    entry_numero_endereco_instalacao.delete(0, 'end')
    entry_bairro_instalacao.delete(0, 'end')
    entry_cidade_instalacao.delete(0, 'end')
    entry_uf_instalacao.delete(0, 'end')
    entry_UC_instalacao.delete(0, 'end')
    entry_concessionaria.delete(0, 'end')

    # Limpa o resultado
    resultado_dados_cliente.configure(text="")


# ==> Função para verificar e atualizar a senha <==
def verificar_e_atualizar_senha():
    nome = entry_seu_nome.get().strip()
    nova_senha = entry_nova_senha.get().strip()
    confirma_nova_senha = entry_confirma_nova_senha.get().strip()

    pergunta1 = combobox_pergunta_recuperacao1.get().strip()
    resposta1 = entry_resposta_recuperacao1.get().strip()
    pergunta2 = combobox_pergunta_recuperacao2.get().strip()
    resposta2 = entry_resposta_recuperacao2.get().strip()
    pergunta3 = combobox_pergunta_recuperacao3.get().strip()
    resposta3 = entry_resposta_recuperacao3.get().strip()

    # Verifica se os campos estão preenchidos
    if not nome or not nova_senha or not confirma_nova_senha or \
       pergunta1 == "Selecione" or not resposta1 or \
       pergunta2 == "Selecione" or not resposta2 or \
       pergunta3 == "Selecione" or not resposta3:
        resultado_recuperacao.configure(text="Preencha todos os campos!", text_color="red")
        return

    if nova_senha != confirma_nova_senha:
        resultado_recuperacao.configure(text="As senhas não coincidem!", text_color="red")
        return

    try:
        # Carrega o Excel
        arquivo_excel = "usuarios_registrados.xlsx"
        df = pd.read_excel(arquivo_excel)

        # Converte a coluna "Senha" para texto, se necessário
        if "Senha" in df.columns:
            df["Senha"] = df["Senha"].astype(str)

        # Verifica se o nome existe
        usuario = df[df["Nome"] == nome]
        if usuario.empty:
            resultado_recuperacao.configure(text="Usuário não encontrado!", text_color="red")
            return

        # Verifica as respostas das perguntas de segurança
        if (usuario["Pergunta 1"].values[0] == pergunta1 and
            usuario["Resposta 1"].values[0].strip().lower() == resposta1.strip().lower() and
            usuario["Pergunta 2"].values[0] == pergunta2 and
            usuario["Resposta 2"].values[0].strip().lower() == resposta2.strip().lower() and
            usuario["Pergunta 3"].values[0] == pergunta3 and
            usuario["Resposta 3"].values[0].strip().lower() == resposta3.strip().lower()):
            
            # Atualiza a senha
            df.loc[df["Nome"] == nome, "Senha"] = nova_senha
            df.to_excel(arquivo_excel, index=False)
            resultado_recuperacao.configure(text="Senha atualizada com sucesso!", text_color="green")
        else:
            resultado_recuperacao.configure(text="Respostas incorretas!", text_color="red")

    except Exception as e:
        resultado_recuperacao.configure(text=f"Erro: {e}", text_color="red")

# ==> Função para alterar o tema <==
def alterar_tema(opcao):
    app._set_appearance_mode(opcao)
  
    # Define as cores de fundo e texto com base no tema
    cor_fundo = "#1a1a1a" if opcao == "dark" else "#f0f0f0"  # Fundo: Escuro ou Claro
    cor_texto = "white" if opcao == "dark" else "black"       # Texto: Branco ou Preto

    # Atualiza a cor de fundo dos frames
    frame_login.configure(fg_color=cor_fundo)
    frame_registro.configure(fg_color=cor_fundo)
    frame_opcoes.configure(fg_color=cor_fundo)
    frame_dados_cliente.configure(fg_color=cor_fundo)
    frame_recuperar_senha.configure(fg_color=cor_fundo)
    frame_termo_aditivo.configure(fg_color=cor_fundo)

    # Atualiza a cor do texto nos elementos da tela de login
    label_login.configure(text_color=cor_texto)
    label_nome_login.configure(text_color=cor_texto)
    label_senha.configure(text_color=cor_texto)
    entry_nome.configure(fg_color = cor_fundo,
                        border_width = 2.5,
                        border_color='#1a1a1a')
    entry_senha_login.configure(fg_color = cor_fundo,
                        border_width = 2.5,
                        border_color='#1a1a1a')
    combobox_tema.configure(fg_color= cor_fundo,    
                        text_color=cor_texto)
    resultado_login.configure(text_color="#FFD700" if opcao == "dark" else "red")  # Mensagem de erro
    label_esqueci_senha.configure(text_color=cor_texto)
    botao_registro.configure(text_color=cor_texto)

    # Atualiza o estado do campo de senha na tela de login
    if checkbox_mostrar_senha_login.get() == 1:  # Se o checkbox estiver marcado
        entry_senha_login.configure(show="")
    else:  # Se o checkbox estiver desmarcado
        entry_senha_login.configure(show="*")

    # Atualiza a cor do texto nos elementos da tela de registro
    label_nome_registro.configure(text_color=cor_texto)
    label_telefone_registro.configure(text_color=cor_texto)
    label_cargo_registro.configure(text_color=cor_texto)
    label_senha_registro.configure(text_color=cor_texto)
    label_confirma_senha_registro.configure(text_color=cor_texto)
    label_pergunta_de_segurança1.configure(text_color=cor_texto)
    label_resposta1.configure(text_color=cor_texto)
    label_pergunta_de_seguranca2.configure(text_color=cor_texto)
    label_resposta2.configure(text_color=cor_texto)
    label_pergunta_de_seguranca3.configure(text_color=cor_texto)
    Label_resposta3.configure(text_color=cor_texto)
    mensagem_registro.configure(text_color="red" if opcao == "dark" else "darkred")  # Mensagem de erro

    # Atualiza a cor do texto nos elementos da tela de recuperação de senha
    labeL_seu_nome.configure(text_color=cor_texto)
    label_nova_senha.configure(text_color=cor_texto)
    _label_confirma_nova_senha.configure(text_color=cor_texto)
    label_pergunta_recuperacao1.configure(text_color=cor_texto)
    entry_resposta_recuperacao1.configure(text_color=cor_texto)
    label_resposta_recuperacao2.configure(text_color=cor_texto)
    entry_resposta_recuperacao2.configure(text_color=cor_texto)
    label_resposta_recuperacao3.configure(text_color=cor_texto)
    entry_resposta_recuperacao3.configure(text_color=cor_texto)
    resultado_recuperacao.configure(text_color="red" if opcao == "dark" else "darkred")  # Mensagem de erro

    # Atualiza a cor do texto dos checkboxes
    checkbox_mostrar_senha_registro.configure(text_color=cor_texto)
    checkbox_mostrar_senha_login.configure(text_color=cor_texto)
    checkbox_mostrar_senha_recuperar.configure(text_color=cor_texto)

# ==> Função do cargo administrador <==

def verificar_permissao_customizar_contrato():
    if cargo_usuario_logado in ['Administrador']:
        mensagem_erro_customizar_contrato.configure(text='')
        mostrar_frame(frame_customizar_contrato)
    else:
        mensagem_erro_customizar_contrato.configure(text='Você precisa ser admnistrador para customizar o contrato!')
        app.after(3000, lambda: mensagem_erro_customizar_contrato.configure(text=''))

# ==> Função para verificar permissões <==
def verificar_permissao_termo_aditivo():
    if cargo_usuario_logado in ["Administrador", "Gerente"]:
        mensagem_erro_termo_aditivo.configure(text="")  # Limpa a mensagem de erro
        mostrar_frame(frame_termo_aditivo)  # Redireciona para a tela de termo aditivo
    else:
        mensagem_erro_termo_aditivo.configure(text="Você não tem permissão para isso. Contate um gerente ou administrador.")
        app.after(3000, lambda: mensagem_erro_termo_aditivo.configure(text=''))

# ==> Função para carregar o contrato <==
def carregar_contrato():
    caminho_arquivo = askopenfilename(filetypes=[("Documentos do Word", "*.docx")])
    if not caminho_arquivo:
        return  # Se o usuário cancelar, não faz nada

    try:
        # Carrega o conteúdo do arquivo .docx
        doc = Document(caminho_arquivo)
        texto = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
        caixa_texto_contrato.delete("1.0", "end")  # Limpa o conteúdo atual
        caixa_texto_contrato.insert("1.0", texto)  # Insere o texto do arquivo
    except Exception as e:
        print(f"Erro ao carregar o arquivo: {e}")

def substituir_por_variavel(variavel):
    try:
        # Obtém o texto selecionado
        texto_selecionado = caixa_texto_contrato.selection_get()
        # Substitui o texto selecionado pela variável
        caixa_texto_contrato.delete("sel.first", "sel.last")
        caixa_texto_contrato.insert("insert", variavel)
    except Exception:
        print("Nenhum texto selecionado!")

# ==> Função para formatar CPF e CNPJ <==
def formatar_cpf_cnpj(event):
    # Obtém o valor atual do campo de entrada
    texto = entry_cpf_cnpj_cliente.get().replace(".", "").replace("-", "").replace("/", "")
    
    # Verifica o valor selecionado na combobox
    tipo = combobox_cpf_cnpj.get()
    
    if tipo == "CPF":
        # Aplica a máscara de CPF (XXX.XXX.XXX-XX)
        if len(texto) > 11:
            texto = texto[:11]  # Limita o tamanho a 11 caracteres
        formatado = f"{texto[:3]}.{texto[3:6]}.{texto[6:9]}-{texto[9:11]}" if len(texto) > 3 else texto
    elif tipo == "CNPJ":
        # Aplica a máscara de CNPJ (XX.XXX.XXX/XXXX-XX)
        if len(texto) > 14:
            texto = texto[:14]  # Limita o tamanho a 14 caracteres
        formatado = f"{texto[:2]}.{texto[2:5]}.{texto[5:8]}/{texto[8:12]}-{texto[12:14]}" if len(texto) > 2 else texto
    else:
        # Se nenhum tipo for selecionado, não aplica máscara
        formatado = texto

    # Atualiza o campo de entrada com o texto formatado
    entry_cpf_cnpj_cliente.delete(0, "end")
    entry_cpf_cnpj_cliente.insert(0, formatado)

# ==> Função para validar os campos de dados do cliente <==
def validar_campos_dados_cliente():
    campos = [
        entry_nome_cliente.get().strip(),
        entry_cpf_cnpj_cliente.get().strip(),
        entry_rua_cliente.get().strip(),
        entry_numero_cliente.get().strip(),
        entry_cidade_cliente.get().strip(),
        entry_uf_cliente.get().strip(),
        entry_cep_cliente.get().strip(),
    ]
    return all(campos)  # Retorna True se todos os campos estiverem preenchidos

# ==> Função para registrar o contrato <==
def registrar_contrato():
    global numero_contrato, caminho_arquivo_contrato

    if not validar_campos_dados_cliente():
        resultado_dados_cliente.configure(text="Preencha todos os campos!", text_color="red")
        return

    # Permite que o cliente escolha o local e o nome do arquivo
    caminho_arquivo_contrato = asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Documentos do Word", "*.docx")],
        title="Salvar Contrato"
    )
    if not caminho_arquivo_contrato:
        return  # Se o cliente cancelar, não faz nada

    try:
        # Cria um novo documento Word
        doc = Document()

        # Adiciona os dados do cliente ao contrato
        doc.add_heading(f"Contrato Nº {numero_contrato}", level=1)
        doc.add_paragraph(f"Nome do Cliente: {entry_nome_cliente.get().strip()}")
        doc.add_paragraph(f"CPF/CNPJ: {entry_cpf_cnpj_cliente.get().strip()}")
        doc.add_paragraph(f"Rua: {entry_rua_cliente.get().strip()},")
        doc.add_paragraph(f'Numero: {entry_numero_cliente.get().strip()}')
        doc.add_paragraph(f"Cidade: {entry_cidade_cliente.get().strip()} - {entry_uf_cliente.get().strip()}")
        doc.add_paragraph(f'UF: {entry_uf_cliente.get().strip()}')
        doc.add_paragraph(f"CEP: {entry_cep_cliente.get().strip()}")
        doc.add_paragraph(f"Quantidade de módulos: {entry_qnts_modulos.get().strip()}") 
        doc.add_paragraph(f"Tipos de módulos: {entry_tipo_modulos.get().strip()}")
        doc.add_paragraph(f"Quantidade de inversores: {entry_inversores.get().strip()}")
        doc.add_paragraph(f'Potencia dos inversores: {entry_potencia_inversores.get().strip()}')
        doc.add_paragraph(f'UC: {entry_UC_instalacao.get().strip()}')
        doc.add_paragraph(f'Concessionária de energia: {entry_concessionaria.get().strip()}')
        doc.add_paragraph(f'Endereço da instalação: {entry_endereco_instalacao.get().strip()}')
        doc.add_paragraph(f'Numero da instalação: {entry_numero_endereco_instalacao.get().strip()}')
        doc.add_paragraph(f'Bairro da instalação: {entry_bairro_instalacao.get().strip()}')
        doc.add_paragraph(f'Cidade da instalação: {entry_cidade_instalacao.get().strip()}')
        doc.add_paragraph(f'UF da instalação: {entry_uf_instalacao.get().strip()}')


        # Salva o contrato no arquivo escolhido
        doc.save(caminho_arquivo_contrato)

        # Incrementa o número do contrato
        numero_contrato += 1

        resultado_dados_cliente.configure(text="Contrato registrado com sucesso!", text_color="green")
    except Exception as e:
        resultado_dados_cliente.configure(text=f"Erro ao registrar contrato: {e}", text_color="red")

# ==> Função para atualizar o contrato <==
def atualizar_contrato():
    if not caminho_arquivo_contrato:
        resultado_dados_cliente.configure(text="Nenhum contrato selecionado!", text_color="red")
        return

    if not validar_campos_dados_cliente():
        resultado_dados_cliente.configure(text="Preencha todos os campos!", text_color="red")
        return

    try:
        # Abre o documento existente
        doc = Document(caminho_arquivo_contrato)

        # Atualiza os dados do contrato
        for paragrafo in doc.paragraphs:
            if "Nome do Cliente:" in paragrafo.text:
                paragrafo.text = f"Nome do Cliente: {entry_nome_cliente.get().strip()}"
            elif "CPF/CNPJ:" in paragrafo.text:
                paragrafo.text = f"CPF/CNPJ: {entry_cpf_cnpj_cliente.get().strip()}"
            elif "Rua:" in paragrafo.text:
                paragrafo.text = f"Rua: {entry_rua_cliente.get().strip()}, {entry_numero_cliente.get().strip()}"
            elif "Cidade:" in paragrafo.text:
                paragrafo.text = f"Cidade: {entry_cidade_cliente.get().strip()} - {entry_uf_cliente.get().strip()}"
            elif "CEP:" in paragrafo.text:
                paragrafo.text = f"CEP: {entry_cep_cliente.get().strip()}"
            elif "Quantidade de módulos:" in paragrafo.text:
                paragrafo.text = f"Quantidade de módulos: {entry_qnts_modulos.get().strip()}"
            elif "Tipos de módulos:" in paragrafo.text:
                paragrafo.text = f"Tipos de módulos: {entry_tipo_modulos.get().strip()}"
            elif "Quantidade de inversores:" in paragrafo.text:
                paragrafo.text = f"Quantidade de inversores: {entry_inversores.get().strip()}"
            elif "Potencia dos inversores:" in paragrafo.text:
                paragrafo.text = f"Potencia dos inversores: {entry_potencia_inversores.get().strip()}"
            elif "Endereço da instalação:" in paragrafo.text:
                paragrafo.text = f"Endereço da instalação: {entry_endereco_instalacao.get().strip()}"
            elif "Numero da instalação:" in paragrafo.text:
                paragrafo.text = f"Numero da instalação: {entry_numero_endereco_instalacao.get().strip()}"
            elif "Bairro da instalação:" in paragrafo.text:
                paragrafo.text = f"Bairro da instalação: {entry_bairro_instalacao.get().strip()}"
            elif "Cidade da instalação:" in paragrafo.text:
                paragrafo.text = f"Cidade da instalação: {entry_cidade_instalacao.get().strip()}"
            elif "UF da instalação:" in paragrafo.text:
                paragrafo.text = f"UF da instalação: {entry_uf_instalacao.get().strip()}"

        # Salva as alterações no mesmo arquivo
        doc.save(caminho_arquivo_contrato)

        resultado_dados_cliente.configure(text="Contrato atualizado com sucesso!", text_color="green")
    except Exception as e:
        resultado_dados_cliente.configure(text=f"Erro ao atualizar contrato: {e}", text_color="red")


# ==> Função de importar contrato ja existente <==
def selecionar_contrato_para_atualizar():
    global caminho_arquivo_contrato

    # Permite que o cliente escolha o arquivo do contrato
    caminho_arquivo_contrato = askopenfilename(
        filetypes=[("Documentos do Word", "*.docx")],
        title="Selecionar Contrato"
    )
    if not caminho_arquivo_contrato:
        return  # Se o cliente cancelar, não faz nada

    try:
        # Atualiza o label com o nome do contrato
        nome_arquivo = caminho_arquivo_contrato.split("/")[-1]  # Extrai o nome do arquivo
        label_nome_contrato.configure(text=f"Contrato: {nome_arquivo}")

        # Carrega os dados do contrato nos campos de entrada
        doc = Document(caminho_arquivo_contrato)
        for paragrafo in doc.paragraphs:
            if "Nome do Cliente:" in paragrafo.text:
                entry_nome_cliente.delete(0, "end")
                entry_nome_cliente.insert(0, paragrafo.text.replace("Nome do Cliente: ", ""))
            elif "CPF/CNPJ:" in paragrafo.text:
                entry_cpf_cnpj_cliente.delete(0, "end")
                entry_cpf_cnpj_cliente.insert(0, paragrafo.text.replace("CPF/CNPJ: ", ""))
            elif "Rua:" in paragrafo.text:
                endereco = paragrafo.text.replace("Rua: ", "").split(", ")
                entry_rua_cliente.delete(0, "end")
                entry_rua_cliente.insert(0, endereco[0])
                if len(endereco) > 1:
                    entry_numero_cliente.delete(0, "end")
                    entry_numero_cliente.insert(0, endereco[1])
            elif "Cidade:" in paragrafo.text:
                cidade_uf = paragrafo.text.replace("Cidade: ", "").split(" - ")
                entry_cidade_cliente.delete(0, "end")
                entry_cidade_cliente.insert(0, cidade_uf[0])
                if len(cidade_uf) > 1:
                    entry_uf_cliente.delete(0, "end")
                    entry_uf_cliente.insert(0, cidade_uf[1])
            elif "CEP:" in paragrafo.text:
                entry_cep_cliente.delete(0, "end")
                entry_cep_cliente.insert(0, paragrafo.text.replace("CEP: ", ""))
            elif "Quantidade de módulos:" in paragrafo.text:
                entry_qnts_modulos.delete(0, "end")
                entry_qnts_modulos.insert(0, paragrafo.text.replace("Quantidade de módulos: ", ""))
            elif "Tipos de módulos:" in paragrafo.text:
                entry_tipo_modulos.delete(0, "end")
                entry_tipo_modulos.insert(0, paragrafo.text.replace("Tipos de módulos: ", ""))
            elif "Quantidade de inversores:" in paragrafo.text:
                entry_inversores.delete(0, "end")
                entry_inversores.insert(0, paragrafo.text.replace("Quantidade de inversores: ", ""))
            elif "Potencia dos inversores:" in paragrafo.text:
                entry_potencia_inversores.delete(0, "end")
                entry_potencia_inversores.insert(0, paragrafo.text.replace("Potencia dos inversores: ", ""))
            elif "Endereço da instalação:" in paragrafo.text:
                entry_endereco_instalacao.delete(0, "end")
                entry_endereco_instalacao.insert(0, paragrafo.text.replace("Endereço da instalação: ", ""))
            elif "Numero da instalação:" in paragrafo.text:
                entry_numero_endereco_instalacao.delete(0, "end")
                entry_numero_endereco_instalacao.insert(0, paragrafo.text.replace("Numero da instalação: ", ""))
            elif "Bairro da instalação:" in paragrafo.text:
                entry_bairro_instalacao.delete(0, "end")
                entry_bairro_instalacao.insert(0, paragrafo.text.replace("Bairro da instalação: ", ""))
            elif "Cidade da instalação:" in paragrafo.text:
                entry_cidade_instalacao.delete(0, "end")
                entry_cidade_instalacao.insert(0, paragrafo.text.replace("Cidade da instalação: ", ""))
            elif "UF da instalação:" in paragrafo.text:
                entry_uf_instalacao.delete(0, "end")
                entry_uf_instalacao.insert(0, paragrafo.text.replace("UF da instalação: ", ""))
    except Exception as e:
        resultado_dados_cliente.configure(text=f"Erro ao carregar contrato: {e}", text_color="red")

# --------------------------------------------------------------------------------
# ===> Frames (Telas) <===
frame_login = CTkFrame(app, fg_color='transparent')
frame_login.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_registro = CTkFrame(app, fg_color='#1a1a1a')  # tela futura
frame_registro.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_recuperar_senha = CTkFrame(app, fg_color='#1a1a1a')  # tela futura
frame_recuperar_senha.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_opcoes = CTkFrame(app, fg_color='#1a1a1a')  # tela opcoes
frame_opcoes.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_dados_cliente = CTkFrame(app, fg_color='#1a1a1a')  # tela futura
frame_dados_cliente.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_dados_financeiros = CTkFrame(app, fg_color='#1a1a1a')  # tela futura
frame_dados_financeiros.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_termo_aditivo = CTkFrame(app, fg_color='#1a1a1a')  # Tela de termo aditivo
frame_termo_aditivo.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)

frame_customizar_contrato = CTkFrame(app, fg_color='#1a1a1a')  # Tela de customização de contrato
frame_customizar_contrato.place(relx=0.5, rely=0.5, anchor='center', relwidth=1, relheight=1)


# --------------------------------------------------------------------------------
# ===> Variaveis Globais <===
# ==> Variáveis para armazenar os dados do usuário <==
# Variável global para armazenar o cargo do usuário logado
cargo_usuario_logado = None
numero_contrato = 1  # Número inicial do contrato
caminho_arquivo_contrato = None # Caminho do arquivo do contrato


# --------------------------------------------------------------------------------
# ===> Tela de Login <===
logo = CTkImage(Image.open('1 pagina.png'), size=(640, 900))

decoracao = CTkLabel(frame_login, text='', image=logo, bg_color='transparent')
decoracao.place(relx=0.44, rely=0.5, anchor='e')

# ComboBox para alterar o tema
combobox_tema = CTkComboBox(
    frame_login,
    values=["Escuro", "Claro"],  # Opções de tema
    font=('Source Sans Pro', 15),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#76ABDF',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=150,
    height=30,
    corner_radius=5,
    state="readonly",  # Impede a entrada de texto
    command=lambda opcao: alterar_tema("dark" if opcao == "Escuro" else "light")
)
combobox_tema.place(relx=0.95, rely=0.05, anchor='ne')
combobox_tema.set("Escuro")  # Define o tema padrão como "Escuro"

label_login = CTkLabel(
    frame_login,
    text='Faça seu login',
    font=('Source Sans Pro', 50, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_login.place(relx=0.6, rely=0.35, anchor='w')

label_nome_login = CTkLabel(
    frame_login,
    text='Digite seu nome:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_nome_login.place(relx=0.6, rely=0.45, anchor='w')

entry_nome = CTkEntry(
    frame_login,
    placeholder_text='Nome',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_nome.place(relx=0.6, rely=0.5, anchor='w')

label_senha = CTkLabel(
    frame_login,
    text='Digite sua senha:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_senha.place(relx=0.6, rely=0.55, anchor='w')

entry_senha_login = CTkEntry(
    frame_login,
    placeholder_text='Senha',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    show='*'
)
entry_senha_login.place(relx=0.6, rely=0.6, anchor='w')

checkbox_mostrar_senha_login = CTkCheckBox(
    frame_login,
    text="Mostrar Senha",
    text_color='white',
    font=('Source Sans Pro', 15),
    command=lambda: alternar_visibilidade_senha_login()
)
checkbox_mostrar_senha_login.place(relx=0.6, rely=0.65, anchor='w')

botao_login = CTkButton(
    frame_login,
    text='Logar',
    width=300,
    height=50,
    border_width=2,
    border_color='white',
    fg_color='#2454FF',
    hover_color='#002387',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=login_realizado
)
botao_login.place(relx=0.6, rely=0.7, anchor='w')

resultado_login = CTkLabel(
    frame_login,
    text="",
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    fg_color='transparent'
)
resultado_login.place(relx=0.6, rely=0.75, anchor='w')

label_esqueci_senha = CTkLabel(
    frame_login,
    text='Esqueci minha senha',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    cursor='hand2'
)
label_esqueci_senha.place(relx=0.6, rely=0.84, anchor='w')

label_esqueci_senha.bind('<Button-1>', lambda e: mostrar_frame(frame_recuperar_senha))


botao_registro = CTkLabel(
    frame_login,
    text='Registrar-se',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    cursor='hand2'
)
botao_registro.place(relx=0.6, rely=0.88, anchor='w')
botao_registro.bind('<Button-1>', lambda e: mostrar_frame(frame_registro))


# --------------------------------------------------------------------------------
# ===> Tela de Registro <===
opcoes_topo = CTkFrame(
    frame_registro,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Faça seu cadastro:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.1, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_recuperacao(), limpar_campos_login(), mostrar_frame(frame_login)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')

label_nome_registro = CTkLabel(
    frame_registro,
    text='Digite seu nome:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_nome_registro.place(relx=0.05, rely=0.1, anchor='w')

entry_nome_registro = CTkEntry(
    frame_registro,
    placeholder_text='Nome',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_nome_registro.place(relx=0.05, rely=0.15, anchor='w')

label_telefone_registro = CTkLabel(
    frame_registro,
    text='Digite seu telefone:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_telefone_registro.place(relx=0.05, rely=0.2, anchor='w')

entry_telefone_registro = CTkEntry(
    frame_registro,
    placeholder_text='Telefone',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_telefone_registro.place(relx=0.05, rely=0.25, anchor='w')

label_cargo_registro = CTkLabel(
    frame_registro,
    text='Digite seu cargo:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_cargo_registro.place(relx=0.05, rely=0.3, anchor='w')

combobox_cargo_registro = CTkComboBox(
    frame_registro,
    values=["Administrador", "Gerente", "Funcionário"],  # Opções de cargo
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=300,
    height=40,
    corner_radius=10,
    state="readonly"  # Impede a entrada de texto
)
combobox_cargo_registro.place(relx=0.05, rely=0.35, anchor='w')
combobox_cargo_registro.set("Selecione")  # Valor padrão

label_senha_registro = CTkLabel(
    frame_registro,
    text='Digite sua senha:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_senha_registro.place(relx=0.05, rely=0.4, anchor='w')

entry_senha_registro = CTkEntry(
    frame_registro,
    placeholder_text='Senha',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10,
    show='*'
)
entry_senha_registro.place(relx=0.05, rely=0.45, anchor='w')

checkbox_mostrar_senha_registro = CTkCheckBox(
    frame_registro,
    text="Mostrar Senha",
    text_color='white',
    font=('Source Sans Pro', 15),
    command=alternar_visibilidade_senha_registro
)
checkbox_mostrar_senha_registro.place(relx=0.05, rely=0.6, anchor='w')

label_confirma_senha_registro = CTkLabel(
    frame_registro,
    text='Confirme sua senha:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_confirma_senha_registro.place(relx=0.05, rely=0.5, anchor='w')

entry_confirma_senha_registro = CTkEntry(
    frame_registro,
    placeholder_text='Confirmação de Senha',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10,
    show='*'
)
entry_confirma_senha_registro.place(relx=0.05, rely=0.55, anchor='w')

label_pergunta_de_segurança1 = CTkLabel(
    frame_registro,
    text='Pergunta de segurança 1:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_pergunta_de_segurança1.place(relx=0.5, rely=0.1, anchor='w')

combobox_pergunta1 = CTkComboBox(
    frame_registro,
    values=["Qual é o seu animal de estimação favorito?", "Qual é o nome da sua mãe?", "Qual é a sua cor favorita?"],  # Opções de pergunta
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"  # Impede a entrada de texto
)
combobox_pergunta1.place(relx=0.5, rely=0.15, anchor='w')

combobox_pergunta1.set("Selecione")  # Valor padrão
label_resposta1 = CTkLabel(
    frame_registro,
    text='Resposta:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_resposta1.place(relx=0.5, rely=0.2, anchor='w')

entry_resposta1 = CTkEntry(
    frame_registro,
    placeholder_text='',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_resposta1.place(relx=0.5, rely=0.25, anchor='w')

label_pergunta_de_seguranca2 = CTkLabel(
    frame_registro,
    text='Pergunta de segurança 2:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_pergunta_de_seguranca2.place(relx=0.5, rely=0.3, anchor='w')

combobox_pergunta2 = CTkComboBox(
    frame_registro,
    values=["Qual marca do seu carro?", "Qual é o nome do seu pai?", "Qual é o seu numero da sorte?"],  # Opções de pergunta
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"  # Impede a entrada de texto
)
combobox_pergunta2.place(relx=0.5, rely=0.35, anchor='w')

combobox_pergunta2.set("Selecione")  # Valor padrão
label_resposta2 = CTkLabel(
    frame_registro,
    text='Resposta:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_resposta2.place(relx=0.5, rely=0.4, anchor='w')

entry_resposta2 = CTkEntry(
    frame_registro,
    placeholder_text='',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_resposta2.place(relx=0.5, rely=0.45, anchor='w')

label_pergunta_de_seguranca3 = CTkLabel(
    frame_registro,
    text='Pergunta de segurança 3:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_pergunta_de_seguranca3.place(relx=0.5, rely=0.5, anchor='w')

combobox_pergunta3 = CTkComboBox(
    frame_registro,
    values=["Qual seu ano de nascimento?", "Qual é a sua comida favorita?", "Qual é o seu time do coração?"],  # Opções de pergunta 
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"  # Impede a entrada de texto
)
combobox_pergunta3.place(relx=0.5, rely=0.55, anchor='w')
combobox_pergunta3.set("Selecione")  # Valor padrão

Label_resposta3 = CTkLabel(
    frame_registro,
    text='Resposta:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
Label_resposta3.place(relx=0.5, rely=0.6, anchor='w')

entry_resposta3 = CTkEntry(
    frame_registro,
    placeholder_text='',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_resposta3.place(relx=0.5, rely=0.65, anchor='w')

botao_cadastro_finalizado = CTkButton(
    frame_registro,
    text='Registrar',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=alertas_registrar_usuario
)
botao_cadastro_finalizado.configure(command=alertas_registrar_usuario)
botao_cadastro_finalizado.place(relx=0.05, rely=0.65, anchor='w')

mensagem_registro = CTkLabel(
    frame_registro,
    text="",
    font=('Source Sans Pro', 20, 'bold'),
    fg_color='transparent'
)
mensagem_registro.place(relx=0.05, rely=0.75, anchor='w')


# --------------------------------------------------------------------------------
# ===> Tela de Recuperar Senha <===
opcoes_topo = CTkFrame(
    frame_recuperar_senha,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Recupere sua senha:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.1, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_recuperacao(),limpar_campos_login(), mostrar_frame(frame_login)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')

labeL_seu_nome = CTkLabel(
    frame_recuperar_senha,
    text='Digite seu nome:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
labeL_seu_nome.place(relx=0.05, rely=0.1, anchor='w')

entry_seu_nome = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Nome',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_seu_nome.place(relx=0.05, rely=0.15, anchor='w')

label_nova_senha = CTkLabel(
    frame_recuperar_senha,
    text='Digite sua nova senha:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_nova_senha.place(relx=0.05, rely=0.2, anchor='w')

entry_nova_senha = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Nova Senha',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10,
    show="*"  # Adicionado para mascarar o texto
)
entry_nova_senha.place(relx=0.05, rely=0.25, anchor='w')

_label_confirma_nova_senha = CTkLabel(
    frame_recuperar_senha,
    text='Confirme sua nova senha:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
_label_confirma_nova_senha.place(relx=0.05, rely=0.3, anchor='w')

entry_confirma_nova_senha = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Confirme Nova Senha',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10,
    show='*'
)
entry_confirma_nova_senha.place(relx=0.05, rely=0.35, anchor='w')

checkbox_mostrar_senha_recuperar = CTkCheckBox(
    frame_recuperar_senha,
    text="Mostrar Senha",
    text_color='white',
    font=('Source Sans Pro', 15),
    command=alternar_visibilidade_senha_recuperar  # Chama a nova função
)
checkbox_mostrar_senha_recuperar.place(relx=0.05, rely=0.40, anchor='w')

# Pergunta de segurança 1
label_pergunta_recuperacao1 = CTkLabel(
    frame_recuperar_senha,
    text='Pergunta de segurança 1:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_pergunta_recuperacao1.place(relx=0.5, rely=0.1, anchor='w')

combobox_pergunta_recuperacao1 = CTkComboBox(
    frame_recuperar_senha,
    values=["Qual é o seu animal de estimação favorito?", "Qual é o nome da sua mãe?", "Qual é a sua cor favorita?"],  # Opções de pergunta
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"
)
combobox_pergunta_recuperacao1.place(relx=0.5, rely=0.15, anchor='w')
combobox_pergunta_recuperacao1.set("Selecione")

entry_resposta_recuperacao1 = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Resposta',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_resposta_recuperacao1.place(relx=0.5, rely=0.2, anchor='w')

label_resposta_recuperacao2 = CTkLabel(
    frame_recuperar_senha,
    text='Pergunta de segurança 2:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_resposta_recuperacao2.place(relx=0.5, rely=0.3, anchor='w')

combobox_pergunta_recuperacao2 = CTkComboBox(
    frame_recuperar_senha,
    values=["Qual marca do seu carro?", "Qual é o nome do seu pai?", "Qual é o seu numero da sorte?"],  # Opções de pergunta
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"
)
combobox_pergunta_recuperacao2.place(relx=0.5, rely=0.35, anchor='w')
combobox_pergunta_recuperacao2.set("Selecione")

entry_resposta_recuperacao2 = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Resposta',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_resposta_recuperacao2.place(relx=0.5, rely=0.4, anchor='w')

label_resposta_recuperacao3 = CTkLabel(
    frame_recuperar_senha,
    text='Pergunta de segurança 3:',
    font=('Source Sans Pro', 20),
    text_color='white',
)
label_resposta_recuperacao3.place(relx=0.5, rely=0.5, anchor='w')

combobox_pergunta_recuperacao3 = CTkComboBox(
    frame_recuperar_senha,
    values=["Qual seu ano de nascimento?", "Qual é a sua comida favorita?", "Qual é o seu time do coração?"],  # Opções de pergunta
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=500,
    height=40,
    corner_radius=10,
    state="readonly"
)
combobox_pergunta_recuperacao3.place(relx=0.5, rely=0.55, anchor='w')
combobox_pergunta_recuperacao3.set("Selecione")

entry_resposta_recuperacao3 = CTkEntry(
    frame_recuperar_senha,
    placeholder_text='Resposta',
    width=500,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
)
entry_resposta_recuperacao3.place(relx=0.5, rely=0.6, anchor='w')

botao_registrar_senha = CTkButton(
    frame_recuperar_senha,
    text='Registrar Nova Senha',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=verificar_e_atualizar_senha
)
botao_registrar_senha.place(relx=0.5, rely=0.7, anchor='w')

resultado_recuperacao = CTkLabel(
    frame_recuperar_senha,
    text="",
    font=('Source Sans Pro', 20, 'bold'),
    text_color='red',
    fg_color='transparent'
)
resultado_recuperacao.place(relx=0.5, rely=0.8, anchor='w')


# --------------------------------------------------------------------------------
# ===> Tela de opcoes <===
opcoes_topo = CTkFrame(
    frame_opcoes,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Escolha uma opção abaixo:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.1, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_login(),mensagem_erro_termo_aditivo.configure(text="") ,mostrar_frame(frame_login)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')

frame_botao_novo_contrato = CTkFrame(
    frame_opcoes,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#383838',
    width=200,
    height=200
)
frame_botao_novo_contrato.place(relx=0.25, rely=0.5, anchor='center')

label_novo_contrato = CTkLabel(
    frame_botao_novo_contrato,
    text='Gerar Novo Contrato',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    wraplength=180,
    justify='center',
    text_color='white',
    fg_color='transparent',
    cursor='hand2'

)
label_novo_contrato.place(relx=0.5, rely=0.5, anchor='center')
label_novo_contrato.bind('<Button-1>', lambda e: mostrar_frame(frame_dados_cliente))

frame_botao_termo_aditivo = CTkFrame(
    frame_opcoes,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#383838',
    width=200,
    height=200
)
frame_botao_termo_aditivo.place(relx=0.5, rely=0.5, anchor='center')

# Mensagem de erro para usuários sem permissão
mensagem_erro_termo_aditivo = CTkLabel(
    frame_opcoes,
    text="",
    font=('Source Sans Pro', 20, 'bold'),
    text_color='red',
    fg_color='transparent'
)
mensagem_erro_termo_aditivo.place(relx=0.5, rely=0.65, anchor='center')

# Botão para gerar termo aditivo
label_termo_aditivo = CTkLabel(
    frame_botao_termo_aditivo,
    text='Gerar Termo Aditivo',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    wraplength=180,
    justify='center',
    text_color='white',
    fg_color='transparent',
    cursor='hand2'
)
label_termo_aditivo.place(relx=0.5, rely=0.5, anchor='center')

# Verifica o cargo antes de redirecionar
label_termo_aditivo.bind('<Button-1>', lambda e: verificar_permissao_termo_aditivo())

frame_botao_customizar_contrato = CTkFrame(
    frame_opcoes,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#383838',
    width=200,
    height=200
)
frame_botao_customizar_contrato.place(relx=0.75, rely=0.5, anchor='center')

# Mensagem de erro para usuários sem permissão
mensagem_erro_customizar_contrato = CTkLabel(
    frame_opcoes,
    text="",
    font=('Source Sans Pro', 20, 'bold'),
    text_color='red',
    fg_color='transparent'
)
mensagem_erro_customizar_contrato.place(relx=0.5, rely=0.65, anchor='center')

label_customizar_contrato = CTkLabel(
    frame_botao_customizar_contrato,
    text='Customizar Contrato',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    wraplength=180,
    justify='center',
    text_color='white',
    fg_color='transparent',
    cursor='hand2'
)
label_customizar_contrato.place(relx=0.5, rely=0.5, anchor='center')
label_customizar_contrato.bind('<Button-1>', lambda e: verificar_permissao_customizar_contrato())


# --------------------------------------------------------------------------------
# ===> Tela de Dados do Cliente <===
opcoes_topo = CTkFrame(
    frame_dados_cliente,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Insira abaixo os dados do cliente:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.12, rely=0.5, anchor='center')

label_opcoes_topo2 = CTkLabel(
    opcoes_topo,
    text='Insira abaixo os dados comerciais:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo2.place(relx=0.5, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_dados_cliente(), mostrar_frame(frame_opcoes)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')

frame_dados_pessoais = CTkFrame(
    frame_dados_cliente,
    width=700,
    height=600,
    corner_radius=20,
    border_width=2,
    border_color='white',
    bg_color='transparent',
    fg_color='transparent'
)
frame_dados_pessoais.place(relx=0.25, rely=0.48, anchor='center')

frame_dados_negociacao = CTkFrame(
    frame_dados_cliente,
    width=700,
    height=600,
    corner_radius=20,
    border_width=2,
    border_color='white',
    bg_color='transparent',
    fg_color='transparent'
)
frame_dados_negociacao.place(relx=0.75, rely=0.48, anchor='center')


label_nome_cliente = CTkLabel(
    frame_dados_pessoais,
    text='Nome do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_nome_cliente.place(relx=0.05, rely=0.03, anchor='w')

entry_nome_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='Nome do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_nome_cliente.place(relx=0.05, rely=0.08, anchor='w')

combobox_cpf_cnpj = CTkComboBox(
    frame_dados_pessoais,
    values=["CPF", "CNPJ"],  # Opções de CPF ou CNPJ
    font=('Source Sans Pro', 20),
    fg_color='#383838',
    text_color='white',
    dropdown_fg_color='#383838',
    dropdown_text_color='white',
    dropdown_font=('Source Sans Pro', 15),
    width=300,
    height=40,
    corner_radius=10,
    state="readonly"  # Impede a entrada de texto
)
combobox_cpf_cnpj.place(relx=0.05, rely=0.16, anchor='w')
combobox_cpf_cnpj.set("CPF ou CNPJ")  # Valor padrão

label_cpf_cnpj_cliente = CTkLabel(
    frame_dados_pessoais,
    text='CPF ou CNPJ do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_cpf_cnpj_cliente.place(relx=0.05, rely=0.23, anchor='w')

entry_cpf_cnpj_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='CPF ou CNPJ do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_cpf_cnpj_cliente.place(relx=0.05, rely=0.29, anchor='w')
entry_cpf_cnpj_cliente.bind("<KeyRelease>", formatar_cpf_cnpj)

label_rua_cliente = CTkLabel(
    frame_dados_pessoais,
    text='Rua do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_rua_cliente.place(relx=0.05, rely=0.35, anchor='w')

entry_rua_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='Rua do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_rua_cliente.place(relx=0.05, rely=0.41, anchor='w')

label_numero_cliente = CTkLabel(
    frame_dados_pessoais,
    text='Número do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_numero_cliente.place(relx=0.05, rely=0.47, anchor='w')

entry_numero_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='Número do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_numero_cliente.place(relx=0.05, rely=0.53, anchor='w')

label_cidade_cliente = CTkLabel(
    frame_dados_pessoais,
    text='Cidade do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_cidade_cliente.place(relx=0.05, rely=0.59, anchor='w')

entry_cidade_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='Cidade do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_cidade_cliente.place(relx=0.05, rely=0.65, anchor='w')

label_uf_cliente = CTkLabel(
    frame_dados_pessoais,
    text='UF do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_uf_cliente.place(relx=0.05, rely=0.71, anchor='w')

entry_uf_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='UF do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_uf_cliente.place(relx=0.05, rely=0.77, anchor='w')

label_CEP_cliente = CTkLabel(
    frame_dados_pessoais,
    text='CEP do Cliente:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_CEP_cliente.place(relx=0.05, rely=0.83, anchor='w')

entry_cep_cliente = CTkEntry(
    frame_dados_pessoais,
    placeholder_text='CEP do Cliente',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_cep_cliente.place(relx=0.05, rely=0.89, anchor='w')

label_qnts_modulos = CTkLabel(
    frame_dados_negociacao,
    text='Quantidade de módulos:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_qnts_modulos.place(relx=0.05, rely=0.03, anchor='w')

entry_qnts_modulos = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Quantidade de módulos',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_qnts_modulos.place(relx=0.05, rely=0.09, anchor='w')

label_tipo_modulos = CTkLabel(
    frame_dados_negociacao,
    text='Tipo de módulos:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_tipo_modulos.place(relx=0.05, rely=0.15, anchor='w')

entry_tipo_modulos = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Tipo de módulos',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_tipo_modulos.place(relx=0.05, rely=0.21, anchor='w')

label_inversores = CTkLabel(
    frame_dados_negociacao,
    text='Quantidade de inversores:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_inversores.place(relx=0.05, rely=0.27, anchor='w')

entry_inversores = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Quantidade de inversores',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_inversores.place(relx=0.05, rely=0.33, anchor='w')

label_potencia_inversores = CTkLabel(
    frame_dados_negociacao,
    text='Potência dos inversores:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_potencia_inversores.place(relx=0.05, rely=0.39, anchor='w')

entry_potencia_inversores = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Potência dos inversores',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_potencia_inversores.place(relx=0.05, rely=0.45, anchor='w')

label_endereco_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='Endereço de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_endereco_instalacao.place(relx=0.05, rely=0.51, anchor='w')

entry_endereco_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Endereço de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_endereco_instalacao.place(relx=0.05, rely=0.57, anchor='w')

label_numero_endereco_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='Número de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_numero_endereco_instalacao.place(relx=0.05, rely=0.63, anchor='w')

entry_numero_endereco_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Número de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_numero_endereco_instalacao.place(relx=0.05, rely=0.69, anchor='w')

label_bairro_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='Bairro de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_bairro_instalacao.place(relx=0.05, rely=0.75, anchor='w')

entry_bairro_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Bairro de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_bairro_instalacao.place(relx=0.05, rely=0.81, anchor='w')

label_cidade_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='Cidade de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_cidade_instalacao.place(relx=0.05, rely=0.87, anchor='w')

entry_cidade_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Cidade de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_cidade_instalacao.place(relx=0.05, rely=0.93, anchor='w')

label_uf_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='UF de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_uf_instalacao.place(relx=0.55, rely=0.03, anchor='w')

entry_uf_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='UF de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_uf_instalacao.place(relx=0.55, rely=0.09, anchor='w')

label_UC_instalacao = CTkLabel(
    frame_dados_negociacao,
    text='UC de instalação:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_UC_instalacao.place(relx=0.55, rely=0.15, anchor='w')

entry_UC_instalacao = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='UC de instalação',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_UC_instalacao.place(relx=0.55, rely=0.21, anchor='w')

label_concessionaria = CTkLabel(
    frame_dados_negociacao,
    text='Concessionária:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_concessionaria.place(relx=0.55, rely=0.27, anchor='w')

entry_concessionaria = CTkEntry(
    frame_dados_negociacao,
    placeholder_text='Concessionária',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_concessionaria.place(relx=0.55, rely=0.33, anchor='w')

resultado_dados_cliente = CTkLabel(
    frame_dados_cliente,
    text="",
    font=('Source Sans Pro', 20, 'bold'),
    text_color='red',
    fg_color='transparent'
)
resultado_dados_cliente.place(relx=0.5, rely=0.905, anchor='center')

botao_registrar_contrato = CTkButton(
    frame_dados_cliente,
    text='Registrar Contrato',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=registrar_contrato
)
botao_registrar_contrato.place(relx=0.2, rely=0.96, anchor='center')

botao_selecionar_contrato = CTkButton(
    frame_dados_cliente,
    text='Selecionar Contrato',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=selecionar_contrato_para_atualizar
)
botao_selecionar_contrato.place(relx=0.5, rely=0.86, anchor='center')

frame_numero_contrato = CTkFrame(
    frame_dados_cliente,
    fg_color='transparent',
    width=300,
    height=40,
    border_color='white',
    border_width=2,
    corner_radius=10,
)
frame_numero_contrato.place(relx=0.75, rely=0.86, anchor='center')

label_nome_contrato = CTkLabel(
    frame_numero_contrato,
    text="Nenhum contrato selecionado",
    font=('Source Sans Pro', 15, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_nome_contrato.place(relx=0.5, rely=0.5, anchor='center')

botao_atualizar_contrato = CTkButton(
    frame_dados_cliente,
    text='Atualizar Contrato',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=atualizar_contrato
)
botao_atualizar_contrato.place(relx=0.5, rely=0.96, anchor='center')

botao_proxima_pagina = CTkButton(
    frame_dados_cliente,
    text='Seguir para Dados Financeiros',
    text_color='white',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    command=lambda: mostrar_frame(frame_dados_financeiros),
    font=('Source Sans Pro', 20, 'bold'),
)
botao_proxima_pagina.place(relx=0.8, rely=0.96, anchor='center')


# ---------------------------------------------------------------------------------
# ==> Tela de Dados Financeiros <==
opcoes_topo = CTkFrame(
    frame_dados_financeiros,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Insira abaixo os dados financeiros:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.15, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: mostrar_frame(frame_opcoes)
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')
botao_voltar.bind('<Button-1>', lambda e: mostrar_frame(frame_dados_cliente))

frame_numero_contrato = CTkFrame(
    frame_dados_financeiros,
    fg_color='transparent',
    width=300,
    height=40,
    border_color='white',
    border_width=2,
    corner_radius=10,
)
frame_numero_contrato.place(relx=0.9, rely=0.1, anchor='e')

label_numero_contrato = CTkLabel(
    frame_numero_contrato,
    text=f"Contrato Nº {numero_contrato}",
    font=('Source Sans Pro', 15, 'bold','underline'),
    text_color='white',
    fg_color='transparent'
)
label_numero_contrato.place(relx=0.5, rely=0.5, anchor='center')

label_valor_contrato = CTkLabel(
    frame_dados_financeiros,
    text='Valor do Contrato:',
    font=('Source Sans Pro', 20),
    text_color='white',
    fg_color='transparent'
)
label_valor_contrato.place(relx=0.1, rely=0.08, anchor='w')

entry_valor_contrato = CTkEntry(
    frame_dados_financeiros,
    placeholder_text='Valor do Contrato',
    width=300,
    height=40,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 20),
    border_color='white',
    border_width=2,
    corner_radius=10
)
entry_valor_contrato.place(relx=0.1, rely=0.13, anchor='w')


# --------------------------------------------------------------------------------
# ===> Tela de Termo Aditivo <===
opcoes_topo = CTkFrame(
    frame_termo_aditivo,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Insira abaixo os dados do termo aditivo:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.15, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_recuperacao(), mostrar_frame(frame_opcoes)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')
botao_voltar.bind('<Button-1>', lambda e: mostrar_frame(frame_opcoes))

# Caixa de texto para inserir informações do termo aditivo
caixa_texto_termo_aditivo = CTkTextbox(
    frame_termo_aditivo,
    width=700,
    height=400,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 18),
    border_color='white',
    border_width=2,
    corner_radius=10
)
caixa_texto_termo_aditivo.place(relx=0.5, rely=0.5, anchor='center')

# Botão para salvar ou processar o texto inserido
botao_finalizar_termo_aditivo = CTkButton(
    frame_termo_aditivo,
    text='Finalizar Termo Aditivo',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
)
botao_finalizar_termo_aditivo.place(relx=0.5, rely=0.78, anchor='center')

# --------------------------------------------------------------------------------
# ===> Tela de Customização de Contrato <===
opcoes_topo = CTkFrame(
    frame_customizar_contrato,
    fg_color='#A9A9A9',
    width=1440,
    height=50
)
opcoes_topo.place(relx=0.5, rely=0, anchor='n')

label_opcoes_topo = CTkLabel(
    opcoes_topo,
    text='Carregue e customize seu contrato:',
    font=('Source Sans Pro', 20, 'bold'),
    text_color='white',
    fg_color='transparent'
)
label_opcoes_topo.place(relx=0.15, rely=0.5, anchor='center')

botao_voltar = CTkButton(
    opcoes_topo,
    text='Voltar',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold', 'underline'),
    command=lambda: [limpar_campos_recuperacao(), mostrar_frame(frame_opcoes)]
)
botao_voltar.place(relx=0.95, rely=0.5, anchor='center')
botao_voltar.bind('<Button-1>', lambda e: mostrar_frame(frame_opcoes))


botao_carregar_contrato = CTkButton(
    frame_customizar_contrato,
    text='Carregar Contrato',
    width=300,
    height=50,
    corner_radius=20,
    border_width=2,
    border_color='white',
    fg_color='#E10F0F',
    hover_color='#dc9f46',
    text_color='white',
    font=('Source Sans Pro', 20, 'bold'),
    command=carregar_contrato
)
botao_carregar_contrato.place(relx=0.15, rely=0.2, anchor='center')

frame_botoes = CTkFrame(
    frame_customizar_contrato,
    fg_color='#383838',
    width=300,
    height=500,
    corner_radius=20
)
frame_botoes.place(relx=0.15, rely=0.6, anchor='center')

botoes_variaveis = [
    ("{NOME_CLIENTE}", "Nome do Cliente"),
    ("{ENDERECO_CLIENTE}", "Endereço do Cliente"),
    ("{VALOR_CONTRATO}", "Valor do Contrato"),
    ("{DATA_CONTRATO}", "Data do Contrato")
]

for i, (variavel, descricao) in enumerate(botoes_variaveis):
    CTkButton(
        frame_botoes,
        text=descricao,
        width=250,
        height=50,
        corner_radius=10,
        fg_color='#E10F0F',
        hover_color='#dc9f46',
        text_color='white',
        font=('Source Sans Pro', 15, 'bold'),
        command=lambda v=variavel: substituir_por_variavel(v)
    ).place(relx=0.5, rely=(0.1 + i * 0.2), anchor='center')

caixa_texto_contrato = CTkTextbox(
    frame_customizar_contrato,
    width=700,
    height=600,
    fg_color='#383838',
    text_color='white',
    font=('Source Sans Pro', 18),
    border_color='white',
    border_width=2,
    corner_radius=10
)
caixa_texto_contrato.place(relx=0.65, rely=0.5, anchor='center')

# --------------------------------------------------------------------------------
# ===> Inicia o aplicativo <===
frame_login.tkraise()
app.mainloop()